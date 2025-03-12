import docx
import os

try:
    # 获取当前目录
    current_dir = os.path.dirname(os.path.abspath(__file__))
    # 构建文档路径
    doc_path = os.path.join(current_dir, 'Mingzi Ye_CV.docx')
    
    # 打开文档
    doc = docx.Document(doc_path)
    
    # 提取段落文本
    paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
    
    # 打印提取的文本
    for para in paragraphs:
        print(para)
    
    # 提取表格内容
    tables_data = []
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text for cell in row.cells]
            table_data.append(row_data)
        tables_data.append(table_data)
    
    # 打印表格内容
    if tables_data:
        print('\n--- 表格内容 ---')
        for i, table in enumerate(tables_data):
            print(f'\n表格 {i+1}:')
            for row in table:
                print(' | '.join(row))
    
except Exception as e:
    print(f'发生错误: {e}')